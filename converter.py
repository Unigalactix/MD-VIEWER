
import io
import markdown
from xhtml2pdf import pisa
from docx import Document
from bs4 import BeautifulSoup
import re
import base64
import requests

def convert_to_pdf(markdown_content: str) -> io.BytesIO:
    """
    Converts markdown content to a PDF file in-memory.
    """
    # 1. Convert Markdown to HTML
    html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
    
    # Add some basic CSS for better PDF rendering
    full_html = f"""
    <html>
    <head>
    <style>
        body {{ font-family: Helvetica, sans-serif; font-size: 12pt; }}
        h1 {{ color: #2c3e50; font-size: 24pt; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
        h2 {{ color: #34495e; font-size: 18pt; margin-top: 20px; }}
        h3 {{ color: #7f8c8d; font-size: 14pt; }}
        code {{ font-family: "Courier New", monospace; background-color: #f4f4f4; padding: 2px 5px; }}
        pre {{ background-color: #f8f8f8; padding: 10px; border: 1px solid #ddd; }}
        blockquote {{ border-left: 4px solid #ddd; padding-left: 10px; color: #666; }}
    </style>
    </head>
    <body>
    {html_content}
    </body>
    </html>
    """

    # 2. Convert HTML to PDF using xhtml2pdf
    pdf_file = io.BytesIO()
    pisa_status = pisa.CreatePDF(
        src=io.StringIO(full_html),
        dest=pdf_file
    )

    if pisa_status.err:
        raise Exception("Error processing PDF")

    pdf_file.seek(0)
    return pdf_file

def fetch_mermaid_image(mermaid_code: str) -> io.BytesIO:
    """
    Fetches a PNG image of the Mermaid chart from mermaid.ink.
    """
    try:
        # 1. Base64 encode the graph definition
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")
        
        # 2. Construct URL
        url = f"https://mermaid.ink/img/{base64_string}"
        
        # 3. Fetch Image
        response = requests.get(url)
        if response.status_code == 200:
            return io.BytesIO(response.content)
        else:
            return None
    except Exception:
        return None

def convert_to_docx(markdown_content: str) -> io.BytesIO:
    """
    Converts markdown content to a DOCX file in-memory.
    """
    # 1. Convert Markdown to HTML first as an intermediate step
    html_content = markdown.markdown(markdown_content, extensions=['extra'])
    soup = BeautifulSoup(html_content, 'html.parser')

    # 2. Create a new DOCX document
    doc = Document()

    # 3. Iterate over HTML elements and map them to DOCX
    # This is a basic mapper.
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            doc.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            doc.add_heading(element.get_text(), level=6)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'pre':
            # Handle code blocks / logs / mermaid
            text_content = element.get_text()
            
            # Simple heuristic to detect mermaid code blocks if they are not explicitly marked with classes in parsing
            # (Markdown parsing often leaves them as <pre><code>...</code></pre> but sometimes classes are stripped or vary)
            is_mermaid = text_content.strip().startswith('graph ') or \
                         text_content.strip().startswith('sequenceDiagram') or \
                         text_content.strip().startswith('classDiagram') or \
                         text_content.strip().startswith('stateDiagram') or \
                         text_content.strip().startswith('erDiagram') or \
                         text_content.strip().startswith('gantt') or \
                         text_content.strip().startswith('pie') or \
                         text_content.strip().startswith('flowchart') or \
                         'mermaid' in   element.get('class', [])

            image_data = None
            if is_mermaid:
                image_data = fetch_mermaid_image(text_content)

            if image_data:
                doc.add_picture(image_data)
            else:
                # It's a regular code block / log
                p = doc.add_paragraph()
                lines = text_content.split('\n')
                for i, line in enumerate(lines):
                    run = p.add_run(line)
                    run.font.name = 'Courier New'
                    if i < len(lines) - 1:
                        fun = p.add_run('\n')
                        fun.font.name = 'Courier New'
            
        elif element.name == 'blockquote':
            doc.add_paragraph(element.get_text(), style='Intense Quote')
        elif element.name == 'table':
            # Handle tables
            rows = element.find_all('tr')
            if not rows:
                continue
            
            # Key assumption: Parsing the first row to determine column count
            first_row_cells = rows[0].find_all(['td', 'th'])
            num_cols = len(first_row_cells)
            
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                # Handle cases where row might have fewer cells than columns
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        table.cell(i, j).text = cell.get_text().strip()

    # 4. Save to BytesIO
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file
